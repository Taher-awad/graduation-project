
import docx

doc = docx.Document("Final_EduVR_Report.md.docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")
for i, p in enumerate(doc.paragraphs[:50]):
    print(f"[{i}] {p.text}")
