
import docx
from docx.shared import Inches
import re
import os

# Configuration
SOURCE_DOC = "Final_EduVR_Report.md.docx"
OUTPUT_DOC = "Final_EduVR_Report_Final.docx"
IMAGE_DIR = "documentation/images"

def process_document():
    if not os.path.exists(SOURCE_DOC):
        print(f"Error: Source file {SOURCE_DOC} not found.")
        return

    doc = docx.Document(SOURCE_DOC)
    regex = re.compile(r"<<<(.+?)>>>")
    
    modified_count = 0

    for paragraph in doc.paragraphs:
        if "<<<" in paragraph.text:
            text = paragraph.text
            matches = regex.findall(text)
            
            for match in matches:
                # Match is like "use_cases.png"
                clean_name = match.strip()
                
                # Check absolute or relative
                image_path = os.path.join(IMAGE_DIR, clean_name)
                
                if not os.path.exists(image_path):
                     # Try typical extensions if missing
                     if os.path.exists(image_path + ".png"):
                         image_path += ".png"
                     else:
                        print(f"Warning: Image not found: {image_path}")
                        continue
                
                print(f"Embedding image: {image_path}")
                
                # Clear the text of the paragraph effectively
                paragraph.clear()
                
                # Add the image
                run = paragraph.add_run()
                try:
                    run.add_picture(image_path, width=Inches(6.0))
                    modified_count += 1
                except Exception as e:
                    print(f"Error embedding {image_path}: {e}")
                    # Restore text if failed
                    paragraph.add_run(f"<<<{match}>>> (Failed to embed)")

    doc.save(OUTPUT_DOC)
    print(f"Done. Saved to {OUTPUT_DOC} with {modified_count} images embedded.")

if __name__ == "__main__":
    process_document()
